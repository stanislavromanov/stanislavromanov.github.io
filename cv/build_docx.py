"""
Build cv.docx as a close, editable match of cv.typ / cv.pdf.

Mirrors the Typst layout: A4, 1cm/0.9cm margins, Charter 10pt, centered
section headers, a full-width rule under the header, right-aligned
locations via a right tab stop, and tight spacing tuned to 2 pages.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Charter"
CONTENT_W = Cm(19.0)  # A4 width 21cm - 2 x 1cm side margins

EN = "–"  # en dash, matches the Typst source

doc = Document()

# Page geometry: A4 with the same margins as the Typst source
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(1.0)
sec.right_margin = Cm(1.0)
sec.top_margin = Cm(0.9)
sec.bottom_margin = Cm(0.9)


def set_font(run, font=FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)


# Charter on the Normal style so the whole document defaults to it
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
_n_rpr = normal.element.get_or_add_rPr()
_n_fonts = _n_rpr.get_or_add_rFonts()
for k in ("w:ascii", "w:hAnsi", "w:cs"):
    _n_fonts.set(qn(k), FONT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0


def para(align=None, before=0.0, after=0.0, line=13.6, justify=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    # exact line height to mimic Typst's tight leading
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if justify:
        p.alignment = ALIGN.JUSTIFY
    elif align:
        p.alignment = align
    return p


def run(p, text, bold=False, italic=False, size=10):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    set_font(r)
    return r


def add_hrule(p):
    """Full-width 1pt bottom border, mirroring #line(stroke: 1pt)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")  # eighths of a point -> 1pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def section(title):
    p = para(align=ALIGN.CENTER, before=7, after=2)
    run(p, title, bold=True, italic=True, size=12)


def bullet(segments, after=1.5):
    """segments: list of (text, bold, italic). Justified like the source."""
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(13.6)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = ALIGN.JUSTIFY
    for text, bold, italic in segments:
        run(p, text, bold=bold, italic=italic, size=10)
    return p


def job(company, location, role_dates, tech):
    """Three-line header block in one paragraph so it never splits."""
    p = para(before=6, after=2)
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    run(p, company, bold=True, size=12)
    run(p, "\t", size=12)
    run(p, location, size=10)
    run(p, "", size=10).add_break()
    run(p, role_dates, bold=True, size=10)
    run(p, "", size=10).add_break()
    run(p, tech, bold=True, italic=True, size=10)


# Header
h = para(align=ALIGN.CENTER, before=0, after=2, line=22)
run(h, "Stanislavs Romanovs", bold=True, size=18)

c = para(align=ALIGN.CENTER, before=0, after=2, line=12)
run(
    c,
    "Berlin, Germany; mail@stanislavromanov.com; stanislavromanov.com; "
    "linkedin.com/in/stanislavromanov; +49 174 897 0750",
    size=9,
)

t = para(align=ALIGN.CENTER, before=0, after=4, line=13)
run(
    t,
    "Senior Full Stack Engineer focused on solutions using .NET, Node, "
    "Typescript, JavaScript, Angular and React",
    bold=True,
    size=10,
)
add_hrule(t)

# Top Skills
section("Top Skills")
top_skills = [
    (".NET", "Developing with .NET professionally since 2009. .NET Core since its launch in 2014"),
    ("Node", "Making web services, micro services and tools since Node 0.8 release in 2012"),
    ("Typescript", "Primarily using Typescript as a replacement for JavaScript since its release"),
    ("Angular", "Writing scalable front-end systems with good Lighthouse scores since beta release of Angular in 2015"),
    ("React", "Creating web and mobile apps using React with Next.js and Expo since 2020"),
    ("AWS", "Using services like Lambda, S3, ECS, AppRunner, EC2, RDS and everything in-between since 2014"),
]
for lead, rest in top_skills:
    bullet([(lead, True, False), (f" {EN} {rest}", False, False)])

# Other Skills
section("Other Skills")
other_skills = [
    ("Backend", "JavaScript, TypeScript, Node, C#, Firebase, Supabase, Shopify, Contentful, RESTful APIs, microservices, serverless computing, test-driven development, CI/CD"),
    ("Frontend", "JavaScript, TypeScript, React, Angular, Next.js, Tailwind CSS, Material UI, Liquid, SSR, SSG, performance optimization"),
    ("Databases", "SQL and NoSQL knowledge, horizontal scaling and performance optimizations, all major ORM knowledge"),
    ("Mobile", "React Native, Native Script, Expo, Android/iOS, Xcode, Android Studio"),
    ("AI/ML", "OpenAI APIs, Stable Diffusion, Hugging Face, AI integrations, RAG implementations with Langchain using Python"),
    ("DevOps", "AWS, AWS CDK, Terraform, Serverless Framework, SST, Docker, Jenkins, self-hosted Linux servers"),
    ("Testing", "Cypress, Playwright, Karma, Jest, BrowserStack"),
]
for lead, rest in other_skills:
    bullet([(lead, True, False), (f" {EN} {rest}", False, False)])

# Work Experience
section("Work Experience")

job(
    "LUMAS Art Editions GmbH", "Berlin, Germany",
    "VP of AI & Engineering, 2026 " + EN + " Present",
    "TypeScript, Node, MCP, Anthropic API, AWS, ECS Fargate, EventBridge, Terraform, Postgres, Microsoft 365",
)
for b in [
    "Designed and built a unified intelligence layer over LUMAS's catalog, analytics, and operational data behind a single MCP server, letting non-technical staff run complex multi-source workflows (landing-page generation, cross-source root-cause analysis, restock recommendations) that previously required cross-functional coordination",
    "Built the Terraform-provisioned data pipeline behind it: an EventBridge-scheduled ECS Fargate task streams pg_dump through a Python redactor into a separate RDS mirror, so AI tooling queries real business data without exposing customer or employee PII",
    "Built an AI grooming assistant for product managers: a panel of agents explores the codebase plus Jira and Figma context to pre-answer blocker questions and write implementation-ready acceptance criteria back to the ticket, cutting grooming time and back-and-forth",
    "Rolled out enterprise AI access for every employee via Claude (Anthropic Teams), where the MCP intelligence layer ships pre-installed for all users, with access and permissions managed through a Microsoft 365 app",
    "Authored the company-wide AI strategy: a prioritized roadmap, explicit \"what we will not do\" boundaries, and a framework for data-backed leadership decisions",
]:
    bullet([(b, False, False)])

job(
    "WhiteWall Media GmbH", "Berlin, Germany",
    "Head of Engineering, 2016 " + EN + " 2026",
    "TypeScript, Angular, Node, .NET C#, React, Shopify, AWS, Docker",
)
for b in [
    "Architected and led the rewrite of the customer-facing eCommerce platform on a modern Angular + SSR stack, hitting a 90+ Lighthouse score and cutting engineering effort while speeding up feature delivery",
    "Conceived, championed, and built the Shopify Extension for resellers; it grew from a side initiative into the company's most strategically important product within six months and won the 2026 TIPA Award",
    "Implemented microservices and containerization with Docker to improve scalability and deployment reliability",
    "Built and led the engineering team - hiring, mentoring, and structuring delivery - while remaining the senior technical voice in cross-functional decisions",
]:
    bullet([(b, False, False)])

job(
    "Spotcap (Rocket Internet)", "Berlin, Germany",
    "Senior Software Engineer, 2014 " + EN + " 2016",
    "JavaScript, Typescript, Node, Angular, React, AWS, Docker",
)
for b in [
    "Migrated a legacy Symfony2 monolith to a microservices-based Node architecture",
    "Rebuilt the front-end application, porting it from a legacy Angular version to React",
    "Developed and maintained a financial loan service in Node processing millions of euros monthly",
    "Enforced strict security standards and optimized API performance on multiple microservices",
]:
    bullet([(b, False, False)])

job(
    "Transmessenger", "Toronto, Canada",
    "Senior Backend Developer, 2013 " + EN + " 2014",
    ".NET C#, React, MongoDB, MSSQL",
)
for b in [
    "Designed and built a multi-tenant CRM with customizable features and secure API endpoints using .NET MVC",
    "Implemented role-based access control to handle millions of records and complex queries",
    "Optimized database performance to support high-volume operations",
]:
    bullet([(b, False, False)])

# Freelance / Contract Work
section("Freelance / Contract Work")

job(
    "GastroHero GmbH", "Remote",
    "Senior Full-Stack Engineer, 2026",
    "TypeScript, Node, Shopify (Remix), AWS (Lambda, S3, DynamoDB), SST, React, Klaviyo",
)
for b in [
    "Migrated sensitive order documents (invoices, credit memos, shipping papers) off publicly accessible Shopify CDN URLs into a private S3 bucket, fronted by a Shopify app proxy enforcing token-based, per-order access control that also works for guest checkouts",
    "Built a daily multi-market monitor crawling the DE/AT/FR/CH storefronts for pricing, availability, and delivery inconsistencies, surfaced through a React dashboard with PDF report export",
    "Delivered several AWS Lambda automations (SST): a Dachser customs-document pipeline, a Klaviyo-driven CRM enrichment service with a DynamoDB audit trail, and a scheduled Dropbox-to-email document mailer",
]:
    bullet([(b, False, False)])

job(
    "PricewaterhouseCoopers (PwC)", "Remote",
    "Senior Full-Stack Software Engineer, 2025",
    ".NET C#, React, Azure, MSSQL, MongoDB, Microservices, Micro-Frontends",
)
for b in [
    "Contributed to the NGA (Next Generation Audit) platform on the Core team, building and maintaining an authentication service with OAuth 2.0 and Azure AD integration",
    "Built admin service for user management, role provisioning, and audit logging",
    "Implemented secure API gateways and ensured compliance with enterprise security standards",
]:
    bullet([(b, False, False)])

job(
    "Transport UK", "Remote",
    "Senior Full-Stack Developer, 2025",
    "React, .NET C#, PostgreSQL, Docker, AWS",
)
for b in [
    "Designed and implemented a comprehensive bus fleet management system from the ground up",
    "Built React frontend with real-time vehicle tracking, route scheduling, and driver management",
    "Developed .NET C# backend with RESTful APIs and PostgreSQL for data persistence",
    "Implemented lost mileage calculation module for fleet cost tracking and financial reconciliation",
]:
    bullet([(b, False, False)])

# Other Projects (no bullets, a single justified paragraph)
op = para(before=6, after=2)
op.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
op.paragraph_format.keep_with_next = True
run(op, "Other Projects", bold=True, size=12)
run(op, "\t", size=12)
run(op, "Remote", size=10)
op.add_run().add_break()
run(op, "2011 " + EN + " 2013", bold=True, size=10)

body = para(before=0, after=0, justify=True)
run(
    body,
    "I have also worked with brands like DKNY, Cisco, SAP, Intel, Autodesk, "
    "and Axis " + EN + " building eCommerce sites, marketing platforms, and "
    "mobile applications. Visit stanislavromanov.com/works for more details.",
    size=10,
)

doc.save("cv.docx")
print("wrote cv.docx")
